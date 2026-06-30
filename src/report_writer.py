"""Business-friendly Chinese report generation utilities."""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .utils import project_path, write_text


STAT_SYMBOLS = {"F", "p", "η²", "ηp²", "M", "SD", "ps"}


def write_business_report_markdown(
    results: dict[str, Any],
    plan: dict[str, Any],
    mean_table: pd.DataFrame,
    figure_path: Path,
    output_path: str | Path = "outputs/business_friendly_report_zh.md",
) -> Path:
    return write_text(output_path, render_report_markdown(results, plan, mean_table, figure_path))


def render_report_markdown(results: dict[str, Any], plan: dict[str, Any], mean_table: pd.DataFrame, figure_path: Path) -> str:
    analysis = plan["analysis"]
    reporting = plan["reporting"]
    outcomes = analysis["outcome_variables"]
    context = _report_context(plan)
    questions = context.get("questions") or _default_questions(plan)
    core_findings = _core_result_paragraphs(results, plan)
    action_recommendations = context.get("action_recommendations") or []

    lines = [
        f"# {reporting['analysis_title']}",
        "",
        "## 1. 核心内容",
        "",
        "### 1.1 报告目标",
        "",
        _trial_overview(plan),
        "",
        "本次分析希望回答两个问题：",
        "",
        *[f"{i}. {question}" for i, question in enumerate(questions, start=1)],
        "",
        "### 1.2 核心结果",
        "",
        *[item for finding in core_findings for item in (finding, "")],
        _format_context_text(context.get("implication_sentence", "核心结果趋势见图1。"), plan),
        "",
        f"**{_figure_title(plan)}**",
        "",
        f"![图 1]({_relative_output_path(figure_path)})",
        "",
        _figure_note_text(results),
        "",
        "### 1.3 行动建议",
        "",
        *[item for recommendation in action_recommendations for item in (_format_context_text(recommendation, plan), "")],
        "",
        "## 2. 详细统计报告",
        "",
        _method_selection_sentence(results, plan),
        "",
        f"**{_table_title(plan)}**",
        "",
        mean_table.to_markdown(index=False),
        "",
        _table_note_text(results),
        "",
    ]
    for i, outcome in enumerate(outcomes):
        prefix = _outcome_paragraph_prefix(i, outcome, plan)
        lines.extend([
            _result_paragraph_markdown(outcome, results["outcomes"][outcome], plan, prefix),
            "",
        ])
    return "\n".join(lines)


def write_business_report_docx(
    results: dict[str, Any],
    plan: dict[str, Any],
    mean_table: pd.DataFrame,
    figure_path: Path,
    output_path: str | Path = "outputs/final_report.docx",
) -> Path:
    doc = Document()
    _setup_document(doc)

    analysis = plan["analysis"]
    reporting = plan["reporting"]
    outcomes = analysis["outcome_variables"]
    context = _report_context(plan)
    questions = context.get("questions") or _default_questions(plan)
    core_findings = _core_result_paragraphs(results, plan)
    action_recommendations = context.get("action_recommendations") or []

    _add_title(doc, reporting["analysis_title"])
    _add_heading(doc, "1. 核心内容", 1)
    _add_heading(doc, "1.1 报告目标", 2)
    _add_body_paragraph(doc, _trial_overview(plan))
    _add_body_paragraph(doc, "本次分析希望回答两个问题：")
    for i, question in enumerate(questions, start=1):
        _add_numbered_paragraph(doc, question, i)
    _add_heading(doc, "1.2 核心结果", 2)
    for finding in core_findings:
        _add_body_paragraph(doc, finding)
    _add_body_paragraph(doc, _format_context_text(context.get("implication_sentence", "核心结果趋势见图1。"), plan))

    _add_caption_block_before_visual(doc, _figure_title(plan))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(project_path(figure_path)), width=Inches(6.45))
    _add_figure_note(doc, results)
    _add_blank_line(doc)

    _add_heading(doc, "1.3 行动建议", 2)
    for recommendation in action_recommendations:
        _add_body_paragraph(doc, _format_context_text(recommendation, plan))

    _add_blank_line(doc)
    _add_heading(doc, "2. 详细统计报告", 1)
    _add_body_paragraph(doc, _method_selection_sentence(results, plan))
    _add_caption_block_before_visual(doc, _table_title(plan))
    _add_mean_table(doc, mean_table)
    _add_table_note(doc, results)
    _add_blank_line(doc)

    for i, outcome in enumerate(outcomes):
        prefix = _outcome_paragraph_prefix(i, outcome, plan)
        _add_result_paragraph_docx(doc, outcome, results["outcomes"][outcome], plan, prefix)

    path = project_path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _format_run(run, size=16, bold=True)


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _format_run(run, size=14 if level == 1 else 12, bold=True, color=RGBColor(31, 77, 120) if level == 1 else None)


def _add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _format_run(run)


def _add_numbered_paragraph(doc: Document, text: str, number: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(22)
    p.paragraph_format.first_line_indent = Pt(-22)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{number}. {text}")
    _format_run(run)


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _format_run(run, size=10, bold=True)


def _add_caption_block_before_visual(doc: Document, text: str) -> None:
    _add_blank_line(doc)
    _add_caption(doc, text)


def _add_blank_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


def _add_figure_note(doc: Document, results: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if _has_significant_pairwise(results):
        _add_text_run(p, "注：点表示个体观测值；箱线图显示中位数和四分位距；半小提琴图显示各组分布。括号表示显著事后比较：* ")
        _add_text_run(p, "p", italic=True)
        _add_text_run(p, " < .05; ** ")
        _add_text_run(p, "p", italic=True)
        _add_text_run(p, " < .01; *** ")
        _add_text_run(p, "p", italic=True)
        _add_text_run(p, " < .001。")
    else:
        _add_text_run(p, "注：点表示个体观测值；箱线图显示中位数和四分位距；半小提琴图显示各组分布。本图未显示显著事后比较括号。")


def _add_table_note(doc: Document, results: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if _has_significant_pairwise(results):
        _add_text_run(p, f"注：数值为均值，括号内为标准差。所有指标均为 1-7 分量表。每一行内，不共享相同字母下标的组别在 {_posthoc_method_label(results)} 事后比较中差异显著，")
        _add_text_run(p, "p", italic=True)
        _add_text_run(p, " < .05。")
    else:
        _add_text_run(p, "注：数值为均值，括号内为标准差。所有指标均为 1-7 分量表。共享相同字母下标表示未发现显著组间差异。")


def _add_mean_table(doc: Document, table_df: pd.DataFrame) -> None:
    rows, cols = table_df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.3)] * cols
    for col_idx, col_name in enumerate(table_df.columns):
        cell = table.cell(0, col_idx)
        cell.width = widths[col_idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_text(cell, str(col_name), bold=True, align_center=col_idx > 0)
    for row_idx, (_, row) in enumerate(table_df.iterrows(), start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if col_idx == 0:
                _set_cell_text(cell, str(value), align_center=False)
            else:
                _set_mean_cell(cell, str(value))
    _set_three_line_borders(table)


def _set_cell_text(cell: Any, text: str, bold: bool = False, align_center: bool = True) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _format_run(run, size=10.5, bold=bold)


def _set_mean_cell(cell: Any, text: str) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left, right = text.split(" ", 1)
    number = "".join(char for char in left if not char.isalpha())
    letters = "".join(char for char in left if char.isalpha())
    _add_text_run(p, number, size=10.5)
    if letters:
        sub = _add_text_run(p, letters, size=8.5)
        sub.font.subscript = True
    _add_text_run(p, " " + right, size=10.5)


def _add_result_paragraph_docx(doc: Document, outcome: str, outcome_result: dict[str, Any], plan: dict[str, Any], prefix: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    parts = _result_paragraph_parts(outcome, outcome_result, plan, prefix)
    for text, style in parts:
        _add_text_run(p, text, italic=style == "stat")


def _result_paragraph_markdown(outcome: str, outcome_result: dict[str, Any], plan: dict[str, Any], prefix: str) -> str:
    rendered = []
    for text, style in _result_paragraph_parts(outcome, outcome_result, plan, prefix):
        rendered.append(f"*{text}*" if style == "stat" else text)
    return "".join(rendered)


def _result_paragraph_parts(outcome: str, outcome_result: dict[str, Any], plan: dict[str, Any], prefix: str) -> list[tuple[str, str | None]]:
    labels = plan["analysis"]["condition_labels_zh"]
    order = plan["analysis"]["condition_order"]
    grouping_label = plan["analysis"]["grouping_variable_label_zh"]
    outcome_label = plan["analysis"]["outcome_labels_zh"][outcome]
    desc = outcome_result["descriptives"]
    omnibus = outcome_result["omnibus"]
    if omnibus["p"] >= plan["reporting"].get("alpha", 0.05):
        return _nonsignificant_result_parts(outcome, outcome_result, plan, prefix)
    baseline_key = _baseline_condition(plan)
    focal_key = _focal_condition(plan)
    baseline = desc[baseline_key]
    focal = desc[focal_key]
    middle_groups = [group for group in order if group not in {baseline_key, focal_key}]
    middle_1 = desc[middle_groups[0]]
    middle_2 = desc[middle_groups[1]]
    pct = _percent_change(outcome, outcome_result, plan)
    negative = _is_negative_outcome(outcome, plan)
    direction = "下降" if negative else "提高"
    focal_pattern = "最低" if negative else "最高"
    comparison_word = "低于" if negative else "高于"
    participant_label = _report_context(plan).get("participant_label", "对象")
    order_sentence = f"{labels[focal_key]}{participant_label}的{outcome_label}{focal_pattern}"
    return [
        (prefix + "单因素方差分析显示，" + grouping_label + "对" + outcome_label + "的影响显著，", None),
        ("F", "stat"),
        (f"({omnibus['df1']:.0f}, {omnibus['df2']:.0f}) = {omnibus['statistic']:.2f}, ", None),
        ("p", "stat"),
        (f" {_format_p_relation(omnibus['p'])}，", None),
        ("η²", "stat"),
        (f" = {_format_no_leading_zero(omnibus['effect_size'], 2)}。事后比较结果显示，不同组别在{outcome_label}上存在明显差异，", None),
        ("ps", "stat"),
        (f" < .05。具体而言，{order_sentence} (", None),
        ("M", "stat"),
        (f" = {focal['mean']:.2f}, ", None),
        ("SD", "stat"),
        (f" = {focal['sd']:.2f})，{labels[middle_groups[1]]}次之 (", None),
        ("M", "stat"),
        (f" = {middle_2['mean']:.2f}, ", None),
        ("SD", "stat"),
        (f" = {middle_2['sd']:.2f})，{labels[middle_groups[0]]}再次之 (", None),
        ("M", "stat"),
        (f" = {middle_1['mean']:.2f}, ", None),
        ("SD", "stat"),
        (f" = {middle_1['sd']:.2f})，均{comparison_word}{labels[baseline_key]}{participant_label} (", None),
        ("M", "stat"),
        (f" = {baseline['mean']:.2f}, ", None),
        ("SD", "stat"),
        (f" = {baseline['sd']:.2f})。相较{labels[baseline_key]}，{labels[focal_key]}使{outcome_label}{direction}约 {pct:.1f}%。", None),
    ]


def _nonsignificant_result_parts(outcome: str, outcome_result: dict[str, Any], plan: dict[str, Any], prefix: str) -> list[tuple[str, str | None]]:
    grouping_label = plan["analysis"]["grouping_variable_label_zh"]
    outcome_label = plan["analysis"]["outcome_labels_zh"][outcome]
    omnibus = outcome_result["omnibus"]
    order = plan["analysis"]["condition_order"]
    labels = plan["analysis"]["condition_labels_zh"]
    desc = outcome_result["descriptives"]
    means = sorted(((group, desc[group]["mean"], desc[group]["sd"]) for group in order), key=lambda item: item[1], reverse=not _is_negative_outcome(outcome, plan))
    best_group, best_mean, best_sd = means[0]
    comparison = "较高" if not _is_negative_outcome(outcome, plan) else "较低"
    return [
        (prefix + "单因素方差分析显示，" + grouping_label + "对" + outcome_label + "的影响未达到统计显著，", None),
        ("F", "stat"),
        (f"({omnibus['df1']:.0f}, {omnibus['df2']:.0f}) = {omnibus['statistic']:.2f}, ", None),
        ("p", "stat"),
        (f" {_format_p_relation(omnibus['p'])}，", None),
        ("η²", "stat"),
        (f" = {_format_no_leading_zero(omnibus['effect_size'], 2)}。描述性结果显示，{labels[best_group]}在{outcome_label}上的均值相对{comparison} (", None),
        ("M", "stat"),
        (f" = {best_mean:.2f}, ", None),
        ("SD", "stat"),
        (f" = {best_sd:.2f})，但当前证据不足以说明不同组别之间存在稳定组间差异，因此不应将该均值差异解读为明确业务效果。", None),
    ]


def _trial_overview(plan: dict[str, Any]) -> str:
    labels = plan["analysis"]["condition_labels_zh"]
    order = plan["analysis"]["condition_order"]
    outcomes = [plan["analysis"]["outcome_labels_zh"][outcome] for outcome in plan["analysis"]["outcome_variables"]]
    overview = _report_context(plan).get("overview_sentence")
    if overview:
        return _format_context_text(overview, plan)
    return f"本次分析对比了{len(order)}个组别在{_join_zh(outcomes)}上的差异，以评估{plan['reporting'].get('decision_context', '当前方案')}。具体而言，分析包含{_join_zh([labels[group] for group in order])}。"


def _percent_change(outcome: str, outcome_result: dict[str, Any], plan: dict[str, Any]) -> float:
    baseline = outcome_result["descriptives"][_baseline_condition(plan)]["mean"]
    focal = outcome_result["descriptives"][_focal_condition(plan)]["mean"]
    if _is_negative_outcome(outcome, plan):
        return (baseline - focal) / baseline * 100
    return (focal - baseline) / baseline * 100


def _core_percentage_sentence(results: dict[str, Any], plan: dict[str, Any]) -> str:
    labels = plan["analysis"]["condition_labels_zh"]
    outcome_labels = plan["analysis"]["outcome_labels_zh"]
    baseline_label = labels[_baseline_condition(plan)]
    focal_label = labels[_focal_condition(plan)]
    increases = []
    decreases = []
    for outcome in plan["analysis"]["outcome_variables"]:
        if results["outcomes"][outcome]["omnibus"]["p"] >= plan["reporting"].get("alpha", 0.05):
            continue
        pct = _percent_change(outcome, results["outcomes"][outcome], plan)
        item = f"{pct:.1f}%的{outcome_labels[outcome]}"
        if _is_negative_outcome(outcome, plan):
            decreases.append(item)
        else:
            increases.append(item)
    parts = []
    if increases:
        parts.append("提升" + "、".join(increases))
    if decreases:
        parts.append("降低" + "、".join(decreases))
    if not parts:
        return "从统计结果看，本轮测试未发现稳定组间差异，因此不宜将描述性均值差异解读为明确业务效果。"
    return f"其中，{focal_label}相较{baseline_label}可以" + "，".join(parts) + "。"


def _core_result_paragraphs(results: dict[str, Any], plan: dict[str, Any]) -> list[str]:
    findings = [_format_context_text(item, plan) for item in _report_context(plan).get("core_findings", [])]
    percent_sentence = _core_percentage_sentence(results, plan)
    if not findings:
        return [percent_sentence]
    findings[-1] = findings[-1] + percent_sentence
    return findings


def _method_selection_sentence(results: dict[str, Any], plan: dict[str, Any]) -> str:
    outcome_count = len(plan["analysis"]["outcome_variables"])
    method = _analysis_method_label(results)
    posthoc = _posthoc_method_label(results)
    return f"在正式分析前，workflow 对数据结构、缺失值、组别样本量、异常值、正态性指标和方差齐性进行了检查。检查结果显示，当前数据基本满足{method}的使用条件。于是，经人工确认后，workflow 对{outcome_count}个结果指标均采用{method}，并使用 {posthoc} 进行事后比较。各组均值、标准差和显著性字母标记见表 1。"


def _analysis_method_label(results: dict[str, Any]) -> str:
    methods = {item["method"] for item in results["outcomes"].values()}
    if methods == {"welch_anova"}:
        return "Welch ANOVA"
    if methods == {"classical_one_way_anova"}:
        return "普通单因素方差分析"
    return "人工确认后的对应方差分析方法"


def _posthoc_method_label(results: dict[str, Any]) -> str:
    methods = {item["posthoc_method"] for item in results["outcomes"].values()}
    if methods == {"games_howell"}:
        return "Games-Howell"
    if methods == {"tukey_hsd"}:
        return "Tukey HSD"
    return "对应的事后比较方法"


def _is_negative_outcome(outcome: str, plan: dict[str, Any]) -> bool:
    return plan["analysis"].get("outcome_directions", {}).get(outcome, "positive") == "negative"


def _baseline_condition(plan: dict[str, Any]) -> str:
    return plan["reporting"].get("baseline_condition", plan["analysis"]["condition_order"][0])


def _focal_condition(plan: dict[str, Any]) -> str:
    return plan["reporting"].get("focal_condition", plan["analysis"]["condition_order"][-1])


def _report_context(plan: dict[str, Any]) -> dict[str, Any]:
    return plan.get("reporting", {}).get("report_context", {})


def _default_questions(plan: dict[str, Any]) -> list[str]:
    grouping = plan["analysis"]["grouping_variable_label_zh"]
    outcomes = _join_zh([plan["analysis"]["outcome_labels_zh"][outcome] for outcome in plan["analysis"]["outcome_variables"]])
    return [f"{grouping}是否会影响{outcomes}？", f"哪个{grouping}水平对应更理想的结果模式？"]


def _format_context_text(text: str, plan: dict[str, Any]) -> str:
    labels = plan["analysis"]["condition_labels_zh"]
    order = plan["analysis"]["condition_order"]
    outcomes = [plan["analysis"]["outcome_labels_zh"][outcome] for outcome in plan["analysis"]["outcome_variables"]]
    mapping = {
        "outcome_list": _join_zh(outcomes),
        "group_labels": "、".join(labels[group] for group in order),
        "grouping_variable_label": plan["analysis"]["grouping_variable_label_zh"],
        "decision_context": plan["reporting"].get("decision_context", ""),
    }
    for index, group in enumerate(order, start=1):
        mapping[f"group_{index}"] = labels[group]
    return text.format(**mapping)


def _outcome_paragraph_prefix(index: int, outcome: str, plan: dict[str, Any]) -> str:
    if index == 0:
        return ""
    if index == 1:
        return "同样地，"
    label = plan["analysis"]["outcome_labels_zh"].get(outcome, outcome)
    return f"在{label}方面，"


def _figure_title(plan: dict[str, Any]) -> str:
    configured_title = _report_context(plan).get("figure_title")
    if configured_title:
        return _format_context_text(configured_title, plan)
    grouping = plan["analysis"]["grouping_variable_label_zh"]
    outcome_count = len(plan["analysis"]["outcome_variables"])
    return f"图 1. 不同{grouping}下{outcome_count}个结果指标的分布与均值趋势"


def _table_title(plan: dict[str, Any]) -> str:
    grouping = plan["analysis"]["grouping_variable_label_zh"]
    return f"表 1. 不同{grouping}下各结果指标的描述性统计"


def _format_p_relation(p: float) -> str:
    if p < 0.001:
        return "< .001"
    return f"= {_format_no_leading_zero(p, 3)}"


def _format_no_leading_zero(value: float, decimals: int) -> str:
    text = f"{value:.{decimals}f}"
    if text.startswith("0."):
        return text[1:]
    if text.startswith("-0."):
        return "-" + text[2:]
    return text


def _has_significant_pairwise(results: dict[str, Any]) -> bool:
    return any(item.get("significant") for outcome in results["outcomes"].values() for item in outcome.get("pairwise", []))


def _figure_note_text(results: dict[str, Any]) -> str:
    if _has_significant_pairwise(results):
        return "注：点表示个体观测值；箱线图显示中位数和四分位距；半小提琴图显示各组分布。括号表示显著事后比较：\\* *p* < .05; \\*\\* *p* < .01; \\*\\*\\* *p* < .001。"
    return "注：点表示个体观测值；箱线图显示中位数和四分位距；半小提琴图显示各组分布。本图未显示显著事后比较括号。"


def _table_note_text(results: dict[str, Any]) -> str:
    if _has_significant_pairwise(results):
        return f"注：数值为均值，括号内为标准差。所有指标均为 1-7 分量表。每一行内，不共享相同字母下标的组别在 {_posthoc_method_label(results)} 事后比较中差异显著，*p* < .05。"
    return "注：数值为均值，括号内为标准差。所有指标均为 1-7 分量表。共享相同字母下标表示未发现显著组间差异。"


def _add_text_run(p: Any, text: str, size: float = 11, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> Any:
    run = p.add_run(text)
    _format_run(run, size=size, bold=bold, italic=italic, color=color)
    return run


def _format_run(run: Any, size: float = 11, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _set_three_line_borders(table: Any) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    for cell in table.rows[0].cells:
        _set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"}, bottom={"val": "single", "sz": "8", "color": "000000"})
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})


def _set_cell_border(cell: Any, **kwargs: Any) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn(f"w:{key}"), str(value))


def _join_zh(items: list[str]) -> str:
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return "、".join(items[:-1]) + "和" + items[-1]


def _relative_output_path(path: Path) -> str:
    try:
        return project_path(path).relative_to(project_path("outputs")).as_posix()
    except ValueError:
        return Path(path).name
